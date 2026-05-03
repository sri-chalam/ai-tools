#!/usr/bin/env python3
"""
YouTube Transcript to Word Document Converter

This script downloads YouTube transcripts and formats them into professional Word documents.
"""

import sys
import re
import os
import time
import random
from urllib.parse import urlparse, parse_qs
from typing import List, Dict, Optional, Tuple
import argparse

try:
    from youtube_transcript_api import YouTubeTranscriptApi
    from youtube_transcript_api.formatters import TextFormatter
except ImportError:
    print("Error: youtube-transcript-api not installed. Run: pip install youtube-transcript-api")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.options import Options
    from selenium.webdriver.common.by import By
    from selenium.webdriver.support.ui import WebDriverWait
    from selenium.webdriver.support import expected_conditions as EC
    import time
except ImportError:
    print("Warning: selenium not installed. Browser fallback will not work.")
    print("Run: pip install selenium webdriver-manager")
    HAS_SELENIUM = False
else:
    HAS_SELENIUM = True
    try:
        from webdriver_manager.chrome import ChromeDriverManager
    except ImportError:
        print("Warning: webdriver-manager not installed. Browser fallback may not work properly.")


class YouTubeTranscriptProcessor:
    def __init__(self):
        self.transcript_data = []
        self.video_title = ""
        self.user_agents = [
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36',
            'Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/118.0.0.0 Safari/537.36'
        ]

    def extract_video_id(self, url: str) -> Optional[str]:
        """Extract video ID from YouTube URL."""
        patterns = [
            r'(?:youtube\.com\/watch\?v=|youtu\.be\/|youtube\.com\/embed\/)([a-zA-Z0-9_-]{11})',
            r'youtube\.com\/v\/([a-zA-Z0-9_-]{11})'
        ]

        for pattern in patterns:
            match = re.search(pattern, url)
            if match:
                return match.group(1)
        return None

    def get_transcript_api(self, video_id: str) -> Optional[List[Dict]]:
        """Get transcript using youtube-transcript-api with retry and user-agent rotation."""
        max_retries = 3
        base_delay = 1  # seconds

        for attempt in range(max_retries):
            try:
                # Try different user agents on retries
                if attempt > 0:
                    # Set a random user agent for this attempt
                    import os
                    os.environ['YOUTUBE_TRANSCRIPT_API_USER_AGENT'] = random.choice(self.user_agents)
                    time.sleep(base_delay * (2 ** attempt))  # Exponential backoff

                transcript = YouTubeTranscriptApi.get_transcript(video_id)

                if self.validate_transcript(transcript):
                    return transcript
                else:
                    print(f"API method returned no valid transcript text (attempt {attempt + 1})")
                    if attempt < max_retries - 1:
                        continue
                    return None

            except Exception as e:
                error_msg = str(e).lower()
                if '403' in error_msg or 'forbidden' in error_msg:
                    print(f"403 Forbidden error on attempt {attempt + 1}: {e}")
                    if attempt < max_retries - 1:
                        print(f"Retrying with different user-agent in {base_delay * (2 ** attempt)} seconds...")
                        continue
                    else:
                        print("All API retry attempts failed with 403 errors")
                        return None
                elif '404' in error_msg or 'not found' in error_msg:
                    print(f"Video not found or transcript unavailable: {e}")
                    return None
                else:
                    print(f"API method failed on attempt {attempt + 1}: {e}")
                    if attempt < max_retries - 1:
                        continue
                    return None

        return None

    def validate_transcript(self, transcript: Optional[List[Dict]]) -> bool:
        """Return True only when transcript contains actual non-empty text segments."""
        if not transcript:
            return False
        for item in transcript:
            if item.get('text') and item['text'].strip():
                return True
        return False

    def check_video_availability(self, url: str) -> bool:
        """Check if the YouTube video is accessible before attempting transcript extraction."""
        try:
            import requests
            headers = {'User-Agent': random.choice(self.user_agents)}
            response = requests.head(url, headers=headers, timeout=10, allow_redirects=True)
            return response.status_code == 200
        except ImportError:
            print("requests not available for video availability check, skipping...")
            return True  # Assume available if we can't check
        except Exception as e:
            print(f"Video availability check failed: {e}")
            return False

    def get_transcript_browser(self, url: str) -> Optional[List[Dict]]:
        """Get transcript using browser automation as fallback with improved error handling."""
        if not HAS_SELENIUM:
            print("Selenium not available for browser fallback.")
            return None

        driver = None
        try:
            chrome_options = Options()
            chrome_options.add_argument("--headless")
            chrome_options.add_argument("--no-sandbox")
            chrome_options.add_argument("--disable-dev-shm-usage")
            chrome_options.add_argument("--disable-blink-features=AutomationControlled")
            chrome_options.add_argument("--disable-web-security")
            chrome_options.add_argument("--disable-features=VizDisplayCompositor")
            chrome_options.add_argument("--disable-extensions")
            chrome_options.add_argument("--disable-plugins")
            chrome_options.add_argument("--disable-gpu")
            chrome_options.add_argument("--no-first-run")
            chrome_options.add_argument("--disable-default-apps")
            chrome_options.add_argument(f"--user-agent={random.choice(self.user_agents)}")
            chrome_options.add_argument("--window-size=1920,1080")

            # Experimental options for better stability
            chrome_options.add_experimental_option("excludeSwitches", ["enable-automation"])
            chrome_options.add_experimental_option('useAutomationExtension', False)

            driver = webdriver.Chrome(ChromeDriverManager().install(), options=chrome_options)

            print("Loading YouTube video page...")
            driver.get(url)

            # Wait for page to load and check if video exists
            WebDriverWait(driver, 15).until(
                EC.presence_of_element_located((By.TAG_NAME, "body"))
            )

            # Check for error messages on the page
            error_selectors = [
                "yt-player-error-message-renderer",
                ".ytp-error",
                "[class*='error']",
                "h1"
            ]

            for selector in error_selectors:
                try:
                    error_element = driver.find_element(By.CSS_SELECTOR, selector)
                    if error_element and "unavailable" in error_element.text.lower():
                        print("Video appears to be unavailable")
                        return None
                except:
                    continue

            # Click on the three dots menu
            try:
                menu_button = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.CSS_SELECTOR, "button[aria-label*='...']"))
                )
                menu_button.click()

                # Click on "Show transcript"
                transcript_button = WebDriverWait(driver, 5).until(
                    EC.element_to_be_clickable((By.XPATH, "//yt-formatted-string[text()='Show transcript']"))
                )
                transcript_button.click()

                # Wait for transcript to load
                WebDriverWait(driver, 5).until(
                    EC.presence_of_element_located((By.CSS_SELECTOR, "ytd-transcript-body-renderer"))
                )

                # Extract transcript segments
                transcript_segments = driver.find_elements(By.CSS_SELECTOR, "ytd-transcript-segment-renderer")

                transcript = []
                for segment in transcript_segments:
                    try:
                        time_element = segment.find_element(By.CSS_SELECTOR, "ytd-transcript-segment-timed-text-renderer")
                        start_time = float(segment.get_attribute("start-offset-sec") or 0)
                        text = time_element.text.strip()

                        if text:
                            transcript.append({
                                'text': text,
                                'start': start_time,
                                'duration': 0  # We'll calculate this later
                            })
                    except:
                        continue

                driver.quit()

                # Calculate durations
                for i in range(len(transcript) - 1):
                    transcript[i]['duration'] = transcript[i + 1]['start'] - transcript[i]['start']
                if transcript:
                    transcript[-1]['duration'] = 5.0  # Default duration for last segment

                if self.validate_transcript(transcript):
                    return transcript
                print("Browser fallback returned no valid transcript text.")
                return None

            except Exception as e:
                print(f"Browser transcript extraction failed: {e}")
                driver.quit()
                return None

        except Exception as e:
            print(f"Browser method failed: {e}")
            return None

    def get_transcript(self, url: str) -> Optional[List[Dict]]:
        """Get transcript using API first, then browser fallback with availability check."""
        video_id = self.extract_video_id(url)
        if not video_id:
            print("Could not extract video ID from URL")
            return None

        print(f"Extracted video ID: {video_id}")

        # Check video availability first
        if not self.check_video_availability(url):
            print("Video appears to be unavailable or inaccessible")
            return None

        # Try API method first
        transcript = self.get_transcript_api(video_id)
        if transcript:
            print("Successfully retrieved transcript using API")
            return transcript

        # Fallback to browser method
        print("API method failed, trying browser fallback...")
        transcript = self.get_transcript_browser(url)
        if transcript:
            print("Successfully retrieved transcript using browser")
            return transcript

        print("All transcript extraction methods failed")
        return None

    def format_timestamp(self, seconds: float) -> str:
        """Format seconds into MM:SS format."""
        minutes = int(seconds // 60)
        seconds = int(seconds % 60)
        return f"{minutes:02d}:{seconds:02d}"

    def generate_title(self, transcript: List[Dict]) -> str:
        """Generate a logical title based on transcript content."""
        if not transcript:
            return "YouTube Transcript"

        # Get first 100 words to analyze
        text = " ".join([item['text'] for item in transcript[:10]])
        words = text.split()[:20]

        # Try to extract meaningful title
        if len(words) > 5:
            # Look for patterns that might indicate a title
            title_candidates = []

            # First sentence or first few words
            first_sentence = text.split('.')[0] if '.' in text else ' '.join(words[:10])
            title_candidates.append(first_sentence)

            # Most common words (simple heuristic)
            common_words = ['the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
            filtered_words = [w for w in words if w.lower() not in common_words and len(w) > 3]
            if filtered_words:
                title_candidates.append(' '.join(filtered_words[:5]))

            # Choose the longest reasonable title
            best_title = max(title_candidates, key=len) if title_candidates else ' '.join(words[:5])
        else:
            best_title = ' '.join(words)

        # Clean up the title
        best_title = re.sub(r'[^\w\s-]', '', best_title)  # Remove special chars except spaces and hyphens
        best_title = ' '.join(best_title.split())  # Normalize spaces

        return best_title[:80] if len(best_title) > 80 else best_title

    def create_document(self, transcript: List[Dict], title: str, output_path: str):
        """Create a formatted Word document from the transcript."""
        doc = Document()

        # Set up styles
        self.setup_styles(doc)

        # Add title
        title_para = doc.add_paragraph()
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add some spacing
        doc.add_paragraph()

        # Process transcript and add content
        self.add_transcript_content(doc, transcript)

        # Save document
        doc.save(output_path)
        print(f"Document saved to: {output_path}")

    def setup_styles(self, doc: Document):
        """Set up custom styles for the document."""
        # Speaker style (bold, colored)
        speaker_style = doc.styles.add_style('Speaker', WD_STYLE_TYPE.PARAGRAPH)
        speaker_style.font.bold = True
        speaker_style.font.color.rgb = RGBColor(0, 100, 0)  # Dark green

        # Timestamp style (italic, smaller)
        timestamp_style = doc.styles.add_style('Timestamp', WD_STYLE_TYPE.PARAGRAPH)
        timestamp_style.font.italic = True
        timestamp_style.font.size = Pt(10)
        timestamp_style.font.color.rgb = RGBColor(128, 128, 128)  # Gray

        # Transcript text style
        transcript_style = doc.styles.add_style('Transcript', WD_STYLE_TYPE.PARAGRAPH)
        transcript_style.font.size = Pt(12)
        transcript_style.paragraph_format.line_spacing = 1.5

    def add_transcript_content(self, doc: Document, transcript: List[Dict]):
        """Add formatted transcript content to the document."""
        current_section = ""
        section_content = []

        for i, item in enumerate(transcript):
            text = item['text'].strip()
            timestamp = self.format_timestamp(item['start'])

            # Add timestamp
            ts_para = doc.add_paragraph(style='Timestamp')
            ts_para.add_run(f"[{timestamp}]")

            # Add transcript text
            text_para = doc.add_paragraph(style='Transcript')
            text_para.add_run(text)

            # Try to identify section breaks (simple heuristic)
            if len(text.split()) > 50:  # Long segments might be section starts
                if current_section and section_content:
                    # Add section heading
                    heading = self.generate_section_heading(section_content)
                    if heading:
                        doc.add_heading(heading, level=2)

                current_section = text[:50] + "..."
                section_content = [text]
            else:
                section_content.append(text)

            # Add some spacing between segments
            if i < len(transcript) - 1:
                doc.add_paragraph()

        # Add final section if any
        if current_section and section_content:
            heading = self.generate_section_heading(section_content)
            if heading:
                doc.add_heading(heading, level=2)

    def generate_section_heading(self, content: List[str]) -> Optional[str]:
        """Generate a heading for a section of transcript."""
        combined_text = " ".join(content)
        sentences = combined_text.split('.')[:2]  # First 1-2 sentences

        if sentences:
            heading = sentences[0].strip()
            # Keep it reasonable length
            if len(heading) > 60:
                heading = heading[:57] + "..."
            return heading

        return None

    def process_video(self, url: str, output_dir: str = ".") -> bool:
        """Main processing function."""
        print(f"Processing YouTube video: {url}")

        # Get transcript
        transcript = self.get_transcript(url)
        if not transcript:
            print("Failed to retrieve transcript")
            return False

        # Generate title
        title = self.generate_title(transcript)
        print(f"Generated title: {title}")

        # Create filename
        filename = self.create_filename(title)
        output_path = os.path.join(output_dir, filename)

        # Create document
        self.create_document(transcript, title, output_path)

        return True

    def create_filename(self, title: str) -> str:
        """Create filename from title."""
        # Convert to title case and replace spaces with dashes
        words = title.split()
        title_case = '-'.join(word.capitalize() for word in words if word)
        return f"{title_case}-YT-Transcript.docx"


def main():
    parser = argparse.ArgumentParser(description='Convert YouTube transcript to Word document')
    parser.add_argument('url', help='YouTube video URL')
    parser.add_argument('-o', '--output', default='.', help='Output directory (default: current directory)')

    args = parser.parse_args()

    processor = YouTubeTranscriptProcessor()
    success = processor.process_video(args.url, args.output)

    if success:
        print("Transcript processing completed successfully!")
        sys.exit(0)
    else:
        print("Transcript processing failed!")
        sys.exit(1)


if __name__ == "__main__":
    main()